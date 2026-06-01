from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "猫神牧场教学手册_学生版.md"
DOCX_PATH = BASE / "猫神牧场教学手册_学生版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DDE3", sz="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_font(run, size=11, bold=False, color="000000", font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_runs_with_code(paragraph, text, size=11):
    index = 0
    while index < len(text):
        start = text.find("`", index)
        if start == -1:
            run = paragraph.add_run(text[index:])
            set_font(run, size=size)
            break

        if start > index:
            run = paragraph.add_run(text[index:start])
            set_font(run, size=size)

        end = text.find("`", start + 1)
        if end == -1:
            run = paragraph.add_run(text[start:])
            set_font(run, size=size)
            break

        run = paragraph.add_run(text[start + 1 : end])
        set_font(run, size=size, font="Consolas")
        index = end + 1


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if len(parsed) < 2:
        return

    header = parsed[0]
    body = parsed[2:] if len(parsed) > 2 else []
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    set_table_borders(table)

    for i, value in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_font(run, size=10.5, bold=True, color="1F1F1F")

    for row_values in body:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            add_runs_with_code(paragraph, value, size=10)
            style_paragraph(paragraph, after=0, line=1.0)

    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_docx():
    doc = Document()
    setup_styles(doc)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    in_code = False
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        if in_code:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            style_paragraph(paragraph, after=2, line=1.0)
            run = paragraph.add_run(line)
            set_font(run, size=9.5, font="Consolas")
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_rows.append(line)
            continue

        flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:].strip())
            set_font(run, size=22, bold=True, color="000000")
            style_paragraph(paragraph, after=10, line=1.0)
        elif line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.add_run(line[3:].strip())
        elif line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            paragraph.add_run(line[4:].strip())
        elif line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            style_paragraph(paragraph, after=8, line=1.25)
            run = paragraph.add_run(line[2:].strip())
            set_font(run, size=10.5, bold=True, color="1F3A5F")
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_runs_with_code(paragraph, line[2:].strip(), size=11)
            style_paragraph(paragraph, after=3, line=1.15)
        elif len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            paragraph = doc.add_paragraph(style="List Number")
            add_runs_with_code(paragraph, line[3:].strip(), size=11)
            style_paragraph(paragraph, after=3, line=1.15)
        else:
            paragraph = doc.add_paragraph()
            add_runs_with_code(paragraph, line.strip(), size=11)
            style_paragraph(paragraph)

    flush_table()
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
